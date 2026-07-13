from docx import Document


def save_interview(text, filename):

    doc = Document()

    for line in text.split("\n"):

        if line.strip():

            if line.endswith(":"):
                doc.add_heading(line, level=2)
            else:
                doc.add_paragraph(line)

    doc.save(filename)