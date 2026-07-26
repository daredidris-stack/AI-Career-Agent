from argparse import ArgumentParser
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
MUTED = RGBColor(82, 93, 108)
BLACK = RGBColor(20, 24, 31)


def set_style_font(style, *, size, color=BLACK, bold=False, italic=False):
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold
    style.font.italic = italic
    fonts = style.element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")


def add_paragraph_style(document, name, *, size, color=BLACK, bold=False):
    styles = document.styles
    style = styles[name] if name in styles else styles.add_style(
        name, WD_STYLE_TYPE.PARAGRAPH
    )
    set_style_font(style, size=size, color=color, bold=bold)
    return style


def configure_styles(document):
    normal = document.styles["Normal"]
    set_style_font(normal, size=11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    name = add_paragraph_style(
        document, "Resume Name", size=24, color=NAVY, bold=True
    )
    name.paragraph_format.space_before = Pt(0)
    name.paragraph_format.space_after = Pt(2)
    name.paragraph_format.keep_with_next = True

    target = add_paragraph_style(
        document, "Resume Target Role", size=11, color=BLUE, bold=True
    )
    target.paragraph_format.space_before = Pt(0)
    target.paragraph_format.space_after = Pt(2)
    target.paragraph_format.keep_with_next = True

    contact = add_paragraph_style(
        document, "Resume Contact", size=9.5, color=MUTED
    )
    contact.paragraph_format.space_before = Pt(0)
    contact.paragraph_format.space_after = Pt(10)
    contact.paragraph_format.keep_with_next = True

    section = add_paragraph_style(
        document, "Resume Section", size=12, color=NAVY, bold=True
    )
    section.paragraph_format.space_before = Pt(10)
    section.paragraph_format.space_after = Pt(4)
    section.paragraph_format.keep_with_next = True

    body = add_paragraph_style(document, "Resume Body", size=10.5)
    body.paragraph_format.space_before = Pt(0)
    body.paragraph_format.space_after = Pt(4)
    body.paragraph_format.line_spacing = 1.15

    experience = add_paragraph_style(
        document, "Resume Experience Header", size=10.5, color=BLACK, bold=True
    )
    experience.paragraph_format.space_before = Pt(4)
    experience.paragraph_format.space_after = Pt(2)
    experience.paragraph_format.keep_with_next = True

    bullet = add_paragraph_style(document, "Resume Bullet", size=10.5)
    bullet.paragraph_format.left_indent = Inches(0.375)
    bullet.paragraph_format.first_line_indent = Inches(-0.188)
    bullet.paragraph_format.space_before = Pt(0)
    bullet.paragraph_format.space_after = Pt(2)
    bullet.paragraph_format.line_spacing = 1.15
    _attach_bullet_numbering(document, bullet)


def _attach_bullet_numbering(document, style):
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(element.get(qn("w:abstractNumId")))
        for element in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(element.get(qn("w:numId")))
        for element in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    number_format = OxmlElement("w:numFmt")
    number_format.set(qn("w:val"), "bullet")
    level.append(number_format)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•")
    level.append(level_text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)
    paragraph_properties = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    paragraph_properties.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    paragraph_properties.append(indent)
    level.append(paragraph_properties)
    abstract.append(level)
    numbering.append(abstract)

    number = OxmlElement("w:num")
    number.set(qn("w:numId"), str(num_id))
    abstract_reference = OxmlElement("w:abstractNumId")
    abstract_reference.set(qn("w:val"), str(abstract_id))
    number.append(abstract_reference)
    numbering.append(number)

    style_properties = style.element.get_or_add_pPr()
    number_properties = OxmlElement("w:numPr")
    level_reference = OxmlElement("w:ilvl")
    level_reference.set(qn("w:val"), "0")
    number_reference = OxmlElement("w:numId")
    number_reference.set(qn("w:val"), str(num_id))
    number_properties.append(level_reference)
    number_properties.append(number_reference)
    style_properties.append(number_properties)


def configure_page(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def add_header(document, full_name, target_role, contact):
    name = document.add_paragraph(style="Resume Name")
    name.alignment = WD_ALIGN_PARAGRAPH.LEFT
    name.add_run(full_name)
    document.add_paragraph(target_role, style="Resume Target Role")
    document.add_paragraph(contact, style="Resume Contact")


def add_section(document, heading, content, *, bullets=False):
    document.add_paragraph(heading.upper(), style="Resume Section")
    if isinstance(content, str):
        content = [content]
    for value in content:
        document.add_paragraph(
            value,
            style="Resume Bullet" if bullets else "Resume Body",
        )


def build_template(destination):
    document = Document()
    configure_page(document)
    configure_styles(document)
    add_header(
        document,
        "{{FULL_NAME}}",
        "{{TARGET_ROLE}}",
        "{{CONTACT_LINE}}",
    )
    add_section(document, "Professional Summary", "{{SUMMARY}}")
    add_section(document, "Core Skills", "{{SKILLS}}")
    add_section(document, "Professional Experience", "{{EXPERIENCE}}")
    add_section(document, "Education", "{{EDUCATION}}")
    add_section(document, "Certifications", "{{CERTIFICATIONS}}")
    add_section(document, "Projects", "{{PROJECTS}}")
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)


def build_reference(destination):
    document = Document()
    configure_page(document)
    configure_styles(document)
    add_header(
        document,
        "Dare Daniel Idris",
        "TARGET ROLE",
        "EMAIL  •  PHONE  •  LOCATION  •  LINKEDIN",
    )
    add_section(
        document,
        "Professional Summary",
        "A concise, truthful summary tailored to the target role and supported by the source résumé.",
    )
    add_section(
        document,
        "Core Skills",
        "AWS • Linux • Python • Networking • Automation • CI/CD",
    )
    document.add_paragraph(
        "PROFESSIONAL EXPERIENCE", style="Resume Section"
    )
    document.add_paragraph(
        "CURRENT ROLE  |  EMPLOYER  |  DATES",
        style="Resume Experience Header",
    )
    for bullet in (
        "Tailored achievement or responsibility supported by the original résumé.",
        "Relevant technology, scope, and outcome written in clear ATS-friendly language.",
        "Additional evidence aligned to the job without inventing experience or metrics.",
    ):
        document.add_paragraph(bullet, style="Resume Bullet")
    add_section(
        document,
        "Education",
        "DEGREE  |  INSTITUTION  |  DATES",
    )
    add_section(
        document,
        "Certifications",
        ["Include only certifications explicitly present in the source résumé."],
        bullets=True,
    )
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)


def main():
    parser = ArgumentParser()
    parser.add_argument("--template", required=True, type=Path)
    parser.add_argument("--reference", required=True, type=Path)
    arguments = parser.parse_args()
    build_template(arguments.template)
    build_reference(arguments.reference)


if __name__ == "__main__":
    main()
