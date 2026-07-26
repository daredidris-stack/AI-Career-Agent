from html import escape
from io import BytesIO
import json
from pathlib import Path
import re

from docx import Document
from docx.oxml.ns import qn
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from backend.services.resume_template_service import resume_template_path


EXPORT_FORMATS = {"txt", "pdf", "docx"}


def export_document(document, export_format: str) -> tuple[bytes, str, str]:
    export_format = export_format.casefold()
    if export_format not in EXPORT_FORMATS:
        raise ValueError("Export format must be txt, pdf, or docx.")

    filename = _safe_filename(document.title)
    if export_format == "txt":
        return (
            _plain_content(document).encode("utf-8"),
            "text/plain; charset=utf-8",
            f"{filename}.txt",
        )
    if export_format == "docx":
        return _docx(document), (
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ), f"{filename}.docx"
    return _pdf(document), "application/pdf", f"{filename}.pdf"


def _docx(document) -> bytes:
    data = _structured_data(document)
    if data is not None:
        template_path = resume_template_path(data.get("template_id"))
        if template_path.exists():
            return _templated_resume_docx(document, data, template_path)

    output = BytesIO()
    word_document = Document()
    word_document.add_heading(document.title, level=0)
    for heading, values in _sections(document):
        if heading:
            word_document.add_heading(heading, level=1)
        for value in values:
            word_document.add_paragraph(value)
    word_document.save(output)
    return output.getvalue()


def _templated_resume_docx(document, data: dict, template_path: Path) -> bytes:
    word_document = Document(str(template_path))
    word_document.core_properties.title = document.title
    word_document.core_properties.author = "NextHire AI"

    _replace_scalar(
        word_document,
        "{{FULL_NAME}}",
        data.get("full_name") or data.get("name"),
    )
    _replace_scalar(
        word_document,
        "{{TARGET_ROLE}}",
        data.get("target_role"),
    )
    _replace_scalar(
        word_document,
        "{{CONTACT_LINE}}",
        data.get("contact_line") or data.get("contact"),
    )
    _replace_scalar(
        word_document,
        "{{SUMMARY}}",
        data.get("summary"),
        remove_heading=True,
    )
    _replace_scalar(
        word_document,
        "{{SKILLS}}",
        " • ".join(_string_values(data.get("skills"))),
        remove_heading=True,
    )
    _replace_experience(word_document, data.get("experience"))
    for token, field in (
        ("{{EDUCATION}}", "education"),
        ("{{CERTIFICATIONS}}", "certifications"),
        ("{{PROJECTS}}", "projects"),
    ):
        _replace_list(word_document, token, data.get(field))

    for paragraph in list(word_document.paragraphs):
        if "{{" in paragraph.text and "}}" in paragraph.text:
            _remove_paragraph(paragraph)

    output = BytesIO()
    word_document.save(output)
    return output.getvalue()


def _pdf(document) -> bytes:
    output = BytesIO()
    styles = getSampleStyleSheet()
    story = [Paragraph(escape(document.title), styles["Title"]), Spacer(1, 18)]
    for heading, values in _sections(document):
        if heading:
            story.append(Paragraph(escape(heading), styles["Heading2"]))
            story.append(Spacer(1, 6))
        for value in values:
            story.append(Paragraph(escape(value) or "&nbsp;", styles["BodyText"]))
            story.append(Spacer(1, 6))
    SimpleDocTemplate(
        output,
        pagesize=LETTER,
        title=document.title,
        author="NextHire AI",
    ).build(story)
    return output.getvalue()


def _safe_filename(value: str) -> str:
    return re.sub(r"[^a-zA-Z0-9_-]+", "-", value).strip("-").lower() \
        or "document"


def _structured_data(document) -> dict | None:
    try:
        data = json.loads(document.content)
    except (TypeError, ValueError):
        return None
    if not isinstance(data, dict) or "skills" not in data:
        return None
    return data


def _replace_scalar(
    word_document,
    token: str,
    value,
    *,
    remove_heading: bool = False,
) -> None:
    text = str(value or "").strip()
    for paragraph in list(word_document.paragraphs):
        if paragraph.text.strip() != token:
            continue
        if text:
            paragraph.text = text
        elif remove_heading:
            _remove_section(paragraph)
        else:
            _remove_paragraph(paragraph)


def _replace_list(word_document, token: str, value) -> None:
    values = _string_values(value)
    for paragraph in list(word_document.paragraphs):
        if paragraph.text.strip() != token:
            continue
        if not values:
            _remove_section(paragraph)
            return
        for item in values:
            paragraph.insert_paragraph_before(item, style="Resume Bullet")
        _remove_paragraph(paragraph)
        return


def _replace_experience(word_document, value) -> None:
    entries = value if isinstance(value, list) else []
    for paragraph in list(word_document.paragraphs):
        if paragraph.text.strip() != "{{EXPERIENCE}}":
            continue
        if not entries:
            _remove_section(paragraph)
            return
        for item in entries:
            if isinstance(item, dict):
                heading = " | ".join(
                    part
                    for part in (
                        str(item.get("role") or "").strip(),
                        str(item.get("company") or "").strip(),
                        str(item.get("dates") or "").strip(),
                    )
                    if part
                )
                if heading:
                    paragraph.insert_paragraph_before(
                        heading,
                        style="Resume Experience Header",
                    )
                for bullet in _string_values(item.get("bullets")):
                    paragraph.insert_paragraph_before(
                        bullet,
                        style="Resume Bullet",
                    )
            else:
                text = str(item).strip()
                if text:
                    paragraph.insert_paragraph_before(
                        text,
                        style="Resume Bullet",
                    )
        _remove_paragraph(paragraph)
        return


def _string_values(value) -> list[str]:
    if isinstance(value, list):
        return [
            str(item).strip()
            for item in value
            if item is not None and str(item).strip()
        ]
    if value:
        return [str(value).strip()]
    return []


def _remove_section(placeholder) -> None:
    previous = placeholder._element.getprevious()
    _remove_paragraph(placeholder)
    if previous is not None and previous.tag == qn("w:p"):
        previous.getparent().remove(previous)


def _remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _sections(document) -> list[tuple[str, list[str]]]:
    data = _structured_data(document)
    if data is None:
        return [("", document.content.split("\n"))]
    sections = []
    for key, label in (
        ("full_name", "Name"),
        ("contact_line", "Contact"),
        ("target_role", "Target Role"),
        ("summary", "Professional Summary"),
        ("skills", "Skills"),
        ("experience", "Experience"),
        ("education", "Education"),
        ("certifications", "Certifications"),
        ("projects", "Projects"),
    ):
        value = data.get(key)
        if key == "experience" and isinstance(value, list):
            values = []
            for item in value:
                if isinstance(item, dict):
                    heading = " | ".join(
                        str(item.get(field) or "").strip()
                        for field in ("role", "company", "dates")
                        if str(item.get(field) or "").strip()
                    )
                    if heading:
                        values.append(heading)
                    values.extend(_string_values(item.get("bullets")))
                elif str(item).strip():
                    values.append(str(item).strip())
        else:
            values = _string_values(value)
        if values:
            sections.append((label, values))
    return sections or [("", [""])]


def _plain_content(document) -> str:
    return "\n\n".join(
        "\n".join(([heading] if heading else []) + values)
        for heading, values in _sections(document)
    )
