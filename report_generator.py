from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet

from docx import Document


def export_resume_report(report):

    doc = SimpleDocTemplate("Resume_Report.pdf")

    styles = getSampleStyleSheet()

    story = []

    story.append(Paragraph("<b>AI Resume Analysis</b>", styles["Title"]))

    story.append(
        Paragraph(
            f"Resume Score: {report['resume_score']}/100",
            styles["Normal"]
        )
    )

    story.append(
        Paragraph(
            f"ATS Score: {report['ats_score']}/100",
            styles["Normal"]
        )
    )

    story.append(Paragraph("<br/><b>Strengths</b>", styles["Heading2"]))

    for item in report["strengths"]:
        story.append(Paragraph(f"• {item}", styles["Normal"]))

    story.append(Paragraph("<br/><b>Improvements</b>", styles["Heading2"]))

    for item in report["improvements"]:
        story.append(Paragraph(f"• {item}", styles["Normal"]))

    doc.build(story)

    print("\n✅ Resume_Report.pdf created successfully!\n")
    
    
def export_resume_report_word(report):
    

    doc = Document()
    ...

    doc = Document()

    doc.add_heading("AI Resume Analysis", level=1)

    doc.add_heading("Resume Scores", level=2)
    doc.add_paragraph(f"Resume Score: {report['resume_score']}/100")
    doc.add_paragraph(f"ATS Score: {report['ats_score']}/100")

    doc.add_heading("Strengths", level=2)

    for item in report["strengths"]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Improvements", level=2)

    for item in report["improvements"]:
        doc.add_paragraph(item, style="List Bullet")

    doc.save("Resume_Report.docx")

    print("\n✅ Resume_Report.docx created successfully!\n")
 