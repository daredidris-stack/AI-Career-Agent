from docx import Document
from docx.shared import Pt


def build_resume(profile, ai_summary, ai_bullets, filename):

    doc = Document()

    title = doc.add_heading(profile["name"], level=0)
    title.runs[0].font.size = Pt(22)

    doc.add_paragraph(
        f'{profile["location"]} | '
        f'{profile["phone"]} | '
        f'{profile["email"]}'
    )

    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(ai_summary)

    doc.add_heading("Technical Skills", level=1)

    for skill in profile["skills"]:
        doc.add_paragraph(skill, style="List Bullet")

    doc.add_heading("Professional Experience", level=1)

    for job in profile["experience"]:

        p = doc.add_paragraph()

        p.add_run(job["role"]).bold = True
        p.add_run(f"\n{job['company']}")
        p.add_run(f"\n{job['start']} - {job['end']}")

        bullets = ai_bullets.get(
            job["role"],
            job["responsibilities"]
        )

        for bullet in bullets:
            doc.add_paragraph(
                bullet,
                style="List Bullet"
            )

    doc.add_heading("Education", level=1)

    edu = profile["education"]

    doc.add_paragraph(
        f"{edu['degree']}\n"
        f"{edu['institution']}\n"
        f"{edu['graduation']}"
    )

    doc.add_heading("Certifications", level=1)

    for cert in profile["certifications"]:
        doc.add_paragraph(cert, style="List Bullet")

    doc.add_heading("Languages", level=1)

    for language, level in profile["languages"].items():
        doc.add_paragraph(
            f"{language}: {level}",
            style="List Bullet"
        )

    doc.save(filename)