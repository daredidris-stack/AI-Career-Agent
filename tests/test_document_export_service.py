import unittest
from io import BytesIO
from types import SimpleNamespace
from zipfile import ZipFile

from docx import Document

from backend.services.document_export_service import export_document


class DocumentExportServiceTests(unittest.TestCase):
    def setUp(self):
        self.document = SimpleNamespace(
            title="Cloud Engineer Resume",
            content="Summary\nAWS and Linux experience",
        )

    def test_exports_valid_pdf(self):
        content, media_type, filename = export_document(self.document, "pdf")

        self.assertTrue(content.startswith(b"%PDF"))
        self.assertEqual(media_type, "application/pdf")
        self.assertEqual(filename, "cloud-engineer-resume.pdf")

    def test_exports_valid_docx(self):
        content, media_type, filename = export_document(self.document, "docx")

        with ZipFile(BytesIO(content)) as archive:
            self.assertIn("word/document.xml", archive.namelist())
        self.assertIn("wordprocessingml.document", media_type)
        self.assertEqual(filename, "cloud-engineer-resume.docx")

    def test_rejects_unknown_export_format(self):
        with self.assertRaises(ValueError):
            export_document(self.document, "exe")

    def test_structured_resume_exports_readable_sections(self):
        self.document.content = (
            '{"summary":"Cloud engineer","skills":["AWS","Linux"],'
            '"experience":["Operated infrastructure"],"education":[]}'
        )

        content, _, _ = export_document(self.document, "txt")

        text = content.decode()
        self.assertIn("Professional Summary\nCloud engineer", text)
        self.assertIn("Skills\nAWS\nLinux", text)

    def test_structured_resume_uses_tailored_word_template(self):
        self.document.content = (
            '{"full_name":"Dare Daniel Idris",'
            '"target_role":"Cloud Engineer",'
            '"contact_line":"dare@example.test | Mexico",'
            '"summary":"Cloud engineer",'
            '"skills":["AWS","Linux"],'
            '"experience":[{"role":"Technician",'
            '"company":"Example Co", "dates":"2022 - Present",'
            '"bullets":["Operated infrastructure"]}],'
            '"education":["Information Technology"],'
            '"certifications":[],"projects":[]}'
        )

        content, _, _ = export_document(self.document, "docx")
        word_document = Document(BytesIO(content))
        rendered_text = "\n".join(
            paragraph.text for paragraph in word_document.paragraphs
        )

        self.assertIn("Dare Daniel Idris", rendered_text)
        self.assertIn("Cloud Engineer", rendered_text)
        self.assertIn("AWS • Linux", rendered_text)
        self.assertIn(
            "Technician | Example Co | 2022 - Present",
            rendered_text,
        )
        self.assertIn("Operated infrastructure", rendered_text)
        self.assertNotIn("{{", rendered_text)
        self.assertNotIn("CERTIFICATIONS", rendered_text)
        self.assertNotIn("PROJECTS", rendered_text)

    def test_structured_resume_uses_selected_template(self):
        for template_id, expected_font in (
            ("ats-professional", "Calibri"),
            ("ats-modern", "Arial"),
            ("ats-classic", "Georgia"),
        ):
            with self.subTest(template_id=template_id):
                self.document.content = (
                    f'{{"template_id":"{template_id}",'
                    '"full_name":"Dare Daniel Idris",'
                    '"target_role":"Operations Leader",'
                    '"contact_line":"Mexico",'
                    '"summary":"Operations professional",'
                    '"skills":["Leadership"],'
                    '"experience":[],"education":[]}'
                )

                content, _, _ = export_document(self.document, "docx")
                word_document = Document(BytesIO(content))

                self.assertEqual(
                    word_document.styles["Resume Name"].font.name,
                    expected_font,
                )

    def test_structured_resume_ignores_null_list_values(self):
        self.document.content = (
            '{"summary":"Cloud engineer",'
            '"skills":["AWS",null,"Linux"],'
            '"experience":[],"education":[]}'
        )

        content, _, _ = export_document(self.document, "txt")

        text = content.decode()
        self.assertIn("Skills\nAWS\nLinux", text)
        self.assertNotIn("None", text)


if __name__ == "__main__":
    unittest.main()
