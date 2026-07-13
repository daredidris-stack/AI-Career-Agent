from ollama import chat
from docx import Document

def tailor_resume(profile, job, analysis):

    prompt = f"""
You are an expert resume writer.

Rewrite this candidate's resume specifically for this job.

Candidate

Name:
{profile['name']}

Education:
{profile['education']}

Experience:
{', '.join(profile['experience'])}

Skills:
{', '.join(profile['skills'])}

Target Roles:
{', '.join(profile['target_roles'])}

Job

Title:
{job['title']}

Company:
{job['company']}

Location:
{job['location']}

Required Skills:
{', '.join(job['skills'])}

AI Job Analysis

Match Score:
{analysis['score']}%

Strengths:
{', '.join(analysis['strengths'])}

Missing Skills:
{', '.join(analysis['missing_skills'])}

Recommendation:
{analysis['recommendation']}

Instructions

You are an elite executive resume writer.

Rewrite this resume specifically for this job.

Rules:

- Never invent experience.
- Never invent certifications.
- Never invent projects.
- Keep everything truthful.
- Optimize for ATS systems.
- Emphasize the candidate's strongest matching skills.
- Rewrite bullet points to sound professional.
- Quantify achievements whenever the provided information allows.
- Prioritize keywords from the job description.
- Make the resume one page.
- Make it suitable for a Cloud Engineer, DevOps Engineer, or AI Infrastructure Engineer role.

Return the resume using EXACTLY this format:

# FULL NAME

Professional Summary

Technical Skills

Professional Experience

Projects

Education

Certifications

Languages

Do not include explanations.

Only output the resume.
"""

    response = chat(
        model="qwen3:8b",
        messages=[
            {
                "role": "user",
                "content": prompt
            }
        ],
    )

    return response.message.content

def save_tailored_resume_docx(text, filename):

    doc = Document()

    title = True

    for line in text.split("\n"):

        line = line.strip()

        if not line:
            continue

        # Resume title (first line)
        if title:
            heading = doc.add_heading(level=0)
            heading.add_run(line).bold = True
            title = False
            continue

        # Section headings
        if line.lower() in [
            "professional summary",
            "technical skills",
            "professional experience",
            "projects",
            "education",
            "certifications",
            "languages"
        ]:
            doc.add_heading(line, level=1)

        # Bullet lists
        elif line.startswith("-"):
            doc.add_paragraph(
                line[1:].strip(),
                style="List Bullet"
            )

        else:
            doc.add_paragraph(line)

    doc.save(filename)